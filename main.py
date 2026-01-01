"""
Smart Document Transformation System
Uses DSPy and FastAPI to learn document transformation patterns from examples

Features:
- Background job processing for large documents
- Real-time progress streaming via Server-Sent Events (SSE)
- Smart chunking with cross-chunk context preservation
- Memory-efficient processing for 100+ page documents
"""

import os
import io
import json
import uuid
import asyncio
from datetime import datetime
from pathlib import Path
from typing import Optional
from concurrent.futures import ThreadPoolExecutor
from dotenv import load_dotenv
from fastapi import FastAPI, File, UploadFile, HTTPException, BackgroundTasks
from fastapi.responses import HTMLResponse, FileResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
import dspy
from pypdf import PdfReader
from docx import Document

# Load environment variables from .env file
load_dotenv()

# ============================================================================
# Storage Configuration
# ============================================================================

STORAGE_DIR = Path("./storage")
UPLOADS_DIR = STORAGE_DIR / "uploads"
OUTPUTS_DIR = STORAGE_DIR / "outputs"
JOBS_DIR = STORAGE_DIR / "jobs"

# Create storage directories
for dir_path in [UPLOADS_DIR, OUTPUTS_DIR, JOBS_DIR]:
    dir_path.mkdir(parents=True, exist_ok=True)

# ============================================================================
# Background Job Processing Infrastructure
# ============================================================================

# Thread pool for running blocking DSPy operations
executor = ThreadPoolExecutor(max_workers=4)

# In-memory job status tracking (for real-time updates)
# Note: In production, use Redis or a database for persistence across restarts
job_status_store: dict[str, dict] = {}


class JobStatus:
    """Job status constants."""
    PENDING = "pending"
    PROCESSING = "processing"
    COMPLETED = "completed"
    FAILED = "failed"


def init_job_status(job_id: str, total_chunks: int = 1) -> dict:
    """Initialize job status tracking."""
    status = {
        "job_id": job_id,
        "status": JobStatus.PENDING,
        "progress": 0,
        "current_chunk": 0,
        "total_chunks": total_chunks,
        "message": "Job queued for processing",
        "started_at": None,
        "completed_at": None,
        "error": None,
        "partial_result": "",
        "analysis": ""
    }
    job_status_store[job_id] = status
    return status


def update_job_status(
    job_id: str,
    status: Optional[str] = None,
    progress: Optional[int] = None,
    current_chunk: Optional[int] = None,
    message: Optional[str] = None,
    partial_result: Optional[str] = None,
    analysis: Optional[str] = None,
    error: Optional[str] = None
):
    """Update job status with new values."""
    if job_id not in job_status_store:
        return

    job = job_status_store[job_id]

    if status is not None:
        job["status"] = status
        if status == JobStatus.PROCESSING and job["started_at"] is None:
            job["started_at"] = datetime.now().isoformat()
        elif status in [JobStatus.COMPLETED, JobStatus.FAILED]:
            job["completed_at"] = datetime.now().isoformat()

    if progress is not None:
        job["progress"] = progress

    if current_chunk is not None:
        job["current_chunk"] = current_chunk

    if message is not None:
        job["message"] = message

    if partial_result is not None:
        job["partial_result"] = partial_result

    if analysis is not None:
        job["analysis"] = analysis

    if error is not None:
        job["error"] = error


def get_job_status(job_id: str) -> Optional[dict]:
    """Get current job status."""
    return job_status_store.get(job_id)


# ============================================================================
# Configuration
# ============================================================================

app = FastAPI(
    title="Smart Document Transformation System",
    description="Learn document transformation patterns from examples and apply them to new documents",
    version="1.0.0"
)

# Mount static files directory
STATIC_DIR = Path(__file__).parent / "static"
TEMPLATES_DIR = Path(__file__).parent / "templates"
app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")

# Configure DSPy with Claude
def configure_dspy():
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise ValueError(
            "ANTHROPIC_API_KEY environment variable is required. "
            "Create a .env file with your API key (see .env.example)"
        )

    # Allow model override via environment variable
    model = os.getenv("ANTHROPIC_MODEL", "anthropic/claude-sonnet-4-20250514")
    # Ensure model has the anthropic/ prefix
    if not model.startswith("anthropic/"):
        model = f"anthropic/{model}"

    lm = dspy.LM(model, api_key=api_key)
    dspy.configure(lm=lm)

# ============================================================================
# Document Extraction Utilities (Optimized for Legal Documents)
# ============================================================================

def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract text content from a PDF file with legal document optimization.

    Features:
    - Page markers for reference
    - Metadata extraction (title, author)
    - Structure preservation
    """
    try:
        reader = PdfReader(io.BytesIO(file_bytes))
        text_parts = []

        # Extract metadata if available
        metadata = reader.metadata
        if metadata:
            meta_lines = []
            if metadata.title:
                meta_lines.append(f"Document Title: {metadata.title}")
            if metadata.author:
                meta_lines.append(f"Author: {metadata.author}")
            if metadata.creation_date:
                meta_lines.append(f"Created: {metadata.creation_date}")
            if meta_lines:
                text_parts.append("[DOCUMENT METADATA]\n" + "\n".join(meta_lines))

        # Extract text page by page with markers
        total_pages = len(reader.pages)
        for page_num, page in enumerate(reader.pages, 1):
            text = page.extract_text()
            if text:
                # Clean up common PDF extraction issues
                # Fix hyphenated words at line breaks
                text = text.replace("-\n", "")
                # Normalize whitespace but preserve paragraph breaks
                lines = text.split("\n")
                cleaned_lines = []
                for line in lines:
                    stripped = line.strip()
                    if stripped:
                        cleaned_lines.append(stripped)
                    elif cleaned_lines and cleaned_lines[-1] != "":
                        cleaned_lines.append("")  # Preserve paragraph break

                page_text = "\n".join(cleaned_lines)
                text_parts.append(f"[PAGE {page_num}/{total_pages}]\n{page_text}")

        return "\n\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract text content from a DOCX file with legal document optimization.

    Features:
    - Heading structure preservation with markers
    - Proper table formatting
    - List handling (numbered and bulleted)
    - Header/footer extraction
    """
    try:
        doc = Document(io.BytesIO(file_bytes))
        text_parts = []

        # Extract core document properties if available
        try:
            core_props = doc.core_properties
            meta_lines = []
            if core_props.title:
                meta_lines.append(f"Document Title: {core_props.title}")
            if core_props.author:
                meta_lines.append(f"Author: {core_props.author}")
            if core_props.created:
                meta_lines.append(f"Created: {core_props.created}")
            if core_props.modified:
                meta_lines.append(f"Modified: {core_props.modified}")
            if meta_lines:
                text_parts.append("[DOCUMENT METADATA]\n" + "\n".join(meta_lines))
        except Exception:
            pass  # Skip if properties unavailable

        # Extract headers from first section if present
        try:
            for section in doc.sections:
                header = section.header
                if header and header.paragraphs:
                    header_text = "\n".join(p.text.strip() for p in header.paragraphs if p.text.strip())
                    if header_text:
                        text_parts.append(f"[HEADER]\n{header_text}")
                break  # Only first section header
        except Exception:
            pass

        # Process document body - paragraphs and tables in order
        text_parts.append("[DOCUMENT BODY]")

        for element in doc.element.body:
            # Handle paragraphs
            if element.tag.endswith('p'):
                for paragraph in doc.paragraphs:
                    if paragraph._element == element:
                        text = paragraph.text.strip()
                        if text:
                            # Detect heading styles
                            style_name = paragraph.style.name if paragraph.style else ""
                            if "Heading" in style_name:
                                # Extract heading level
                                level = "".join(filter(str.isdigit, style_name)) or "1"
                                text_parts.append(f"\n[HEADING {level}] {text}")
                            elif "Title" in style_name:
                                text_parts.append(f"\n[TITLE] {text}")
                            else:
                                # Check for list formatting
                                if paragraph._element.pPr is not None:
                                    numPr = paragraph._element.pPr.numPr
                                    if numPr is not None:
                                        text_parts.append(f"  • {text}")
                                    else:
                                        text_parts.append(text)
                                else:
                                    text_parts.append(text)
                        break

            # Handle tables
            elif element.tag.endswith('tbl'):
                for table in doc.tables:
                    if table._element == element:
                        table_lines = ["[TABLE START]"]
                        for row_idx, row in enumerate(table.rows):
                            cells = []
                            for cell in row.cells:
                                cell_text = cell.text.strip().replace("\n", " ")
                                cells.append(cell_text if cell_text else "-")
                            row_text = " | ".join(cells)
                            if row_idx == 0:
                                table_lines.append(row_text)
                                table_lines.append("-" * len(row_text))  # Header separator
                            else:
                                table_lines.append(row_text)
                        table_lines.append("[TABLE END]")
                        text_parts.append("\n".join(table_lines))
                        break

        # Extract footers
        try:
            for section in doc.sections:
                footer = section.footer
                if footer and footer.paragraphs:
                    footer_text = "\n".join(p.text.strip() for p in footer.paragraphs if p.text.strip())
                    if footer_text:
                        text_parts.append(f"\n[FOOTER]\n{footer_text}")
                break
        except Exception:
            pass

        return "\n\n".join(text_parts)
    except Exception as e:
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Extract text content from a plain text file."""
    try:
        return file_bytes.decode("utf-8")
    except UnicodeDecodeError:
        try:
            return file_bytes.decode("latin-1")
        except Exception as e:
            raise ValueError(f"Failed to decode text file: {str(e)}")


def extract_text(filename: str, file_bytes: bytes) -> str:
    """Router function to extract text based on file extension."""
    filename_lower = filename.lower()

    if filename_lower.endswith(".pdf"):
        return extract_text_from_pdf(file_bytes)
    elif filename_lower.endswith(".docx"):
        return extract_text_from_docx(file_bytes)
    elif filename_lower.endswith((".txt", ".md", ".json", ".xml", ".csv")):
        return extract_text_from_txt(file_bytes)
    else:
        # Try to decode as text for unknown extensions
        try:
            return extract_text_from_txt(file_bytes)
        except ValueError:
            raise ValueError(f"Unsupported file format: {filename}")


# ============================================================================
# Storage Utilities
# ============================================================================

def generate_job_id() -> str:
    """Generate a unique job ID."""
    return f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid.uuid4().hex[:8]}"


def save_uploaded_file(file_bytes: bytes, original_filename: str, job_id: str, file_type: str) -> Path:
    """Save an uploaded file to storage."""
    job_uploads_dir = UPLOADS_DIR / job_id
    job_uploads_dir.mkdir(parents=True, exist_ok=True)

    # Preserve original extension
    ext = Path(original_filename).suffix
    saved_filename = f"{file_type}{ext}"
    file_path = job_uploads_dir / saved_filename

    file_path.write_bytes(file_bytes)
    return file_path


def save_output(job_id: str, transformed_text: str, analysis: str, metadata: dict) -> Path:
    """Save transformation output to storage."""
    job_output_dir = OUTPUTS_DIR / job_id
    job_output_dir.mkdir(parents=True, exist_ok=True)

    # Save transformed document
    output_file = job_output_dir / "transformed_output.txt"
    output_file.write_text(transformed_text, encoding="utf-8")

    # Save analysis
    analysis_file = job_output_dir / "analysis.txt"
    analysis_file.write_text(analysis, encoding="utf-8")

    # Save metadata as JSON
    metadata_file = job_output_dir / "metadata.json"
    metadata_file.write_text(json.dumps(metadata, indent=2), encoding="utf-8")

    return output_file


def save_job_record(job_id: str, job_data: dict) -> Path:
    """Save a job record with all metadata."""
    job_file = JOBS_DIR / f"{job_id}.json"
    job_data["job_id"] = job_id
    job_data["created_at"] = datetime.now().isoformat()
    job_file.write_text(json.dumps(job_data, indent=2), encoding="utf-8")
    return job_file


def list_jobs() -> list[dict]:
    """List all saved transformation jobs."""
    jobs = []
    for job_file in sorted(JOBS_DIR.glob("*.json"), reverse=True):
        try:
            job_data = json.loads(job_file.read_text())
            jobs.append(job_data)
        except (json.JSONDecodeError, IOError):
            continue
    return jobs


def get_job(job_id: str) -> Optional[dict]:
    """Get a specific job by ID."""
    job_file = JOBS_DIR / f"{job_id}.json"
    if job_file.exists():
        return json.loads(job_file.read_text())
    return None


# ============================================================================
# DSPy Transformation Module
# ============================================================================

class DocumentTransformation(dspy.Signature):
    """Analyze a document transformation pattern from an example pair and apply it to a new document."""

    example_input: str = dspy.InputField(
        desc="Original document BEFORE transformation"
    )
    example_output: str = dspy.InputField(
        desc="Document AFTER transformation - the target format/style"
    )
    new_document: str = dspy.InputField(
        desc="New document to transform using the learned pattern"
    )

    transformation_analysis: str = dspy.OutputField(
        desc="Analysis of what transformation was applied: structural changes, content modifications, formatting rules, extraction patterns"
    )
    transformed_document: str = dspy.OutputField(
        desc="The new document transformed following the exact same pattern as the example"
    )


class SmartDocumentTransformer(dspy.Module):
    """DSPy module that learns and applies document transformations."""

    def __init__(self):
        super().__init__()
        self.transform = dspy.ChainOfThought(DocumentTransformation)

    def forward(self, example_input: str, example_output: str, new_document: str):
        result = self.transform(
            example_input=example_input,
            example_output=example_output,
            new_document=new_document
        )
        return result


class MultiExampleTransformation(dspy.Signature):
    """Learn transformation patterns from multiple example pairs for better accuracy."""

    examples_summary: str = dspy.InputField(
        desc="Summary of multiple input/output example pairs showing the transformation pattern"
    )
    new_document: str = dspy.InputField(
        desc="New document to transform using the learned patterns"
    )

    transformation_analysis: str = dspy.OutputField(
        desc="Detailed analysis of the transformation pattern learned from all examples"
    )
    transformed_document: str = dspy.OutputField(
        desc="The new document transformed following the learned pattern"
    )


class MultiExampleTransformer(dspy.Module):
    """DSPy module for learning from multiple example pairs."""

    def __init__(self):
        super().__init__()
        self.transform = dspy.ChainOfThought(MultiExampleTransformation)

    def forward(self, examples_summary: str, new_document: str):
        result = self.transform(
            examples_summary=examples_summary,
            new_document=new_document
        )
        return result


# ============================================================================
# Context-Aware Transformation for Large Documents
# ============================================================================

class ContextAwareTransformation(dspy.Signature):
    """Transform a document chunk while maintaining context from previous chunks."""

    example_input: str = dspy.InputField(
        desc="Original document BEFORE transformation"
    )
    example_output: str = dspy.InputField(
        desc="Document AFTER transformation - the target format/style"
    )
    previous_context: str = dspy.InputField(
        desc="Summary of what was transformed in previous chunks (empty for first chunk)"
    )
    current_chunk: str = dspy.InputField(
        desc="Current chunk of the document to transform"
    )
    chunk_info: str = dspy.InputField(
        desc="Information about chunk position (e.g., 'Chunk 2 of 5')"
    )

    transformation_analysis: str = dspy.OutputField(
        desc="Analysis of transformation applied to this chunk"
    )
    transformed_chunk: str = dspy.OutputField(
        desc="The transformed chunk following the same pattern as the example"
    )
    context_summary: str = dspy.OutputField(
        desc="Brief summary of key elements transformed in this chunk (for next chunk's context)"
    )


class ContextAwareTransformer(dspy.Module):
    """DSPy module that maintains context across chunks for large documents."""

    def __init__(self):
        super().__init__()
        self.transform = dspy.ChainOfThought(ContextAwareTransformation)

    def forward(
        self,
        example_input: str,
        example_output: str,
        previous_context: str,
        current_chunk: str,
        chunk_info: str
    ):
        result = self.transform(
            example_input=example_input,
            example_output=example_output,
            previous_context=previous_context,
            current_chunk=current_chunk,
            chunk_info=chunk_info
        )
        return result


def detect_transformation_type(input_text: str, output_text: str) -> str:
    """Detect the type of transformation based on input/output characteristics."""
    input_len = len(input_text)
    output_len = len(output_text)
    ratio = output_len / input_len if input_len > 0 else 1

    # Check for structured output patterns
    has_json_output = output_text.strip().startswith("{") or output_text.strip().startswith("[")
    has_sections = any(marker in output_text for marker in ["##", "**", "---", "==="])
    has_bullet_points = any(marker in output_text for marker in ["- ", "* ", "• "])

    if ratio < 0.3:
        return "summarization"
    elif ratio > 2.0:
        return "expansion"
    elif has_json_output:
        return "data_extraction"
    elif has_sections or has_bullet_points:
        return "reformatting"
    elif abs(ratio - 1.0) < 0.3:
        return "style_transfer"
    else:
        return "general_transformation"


def chunk_text(text: str, max_chars: int = 15000, overlap: int = 1000) -> list[str]:
    """
    Split text into chunks for processing large documents.

    Optimized for legal documents with:
    - Larger chunks (15k chars ~= 3-4 pages) to maintain context
    - Smart boundary detection (sections, paragraphs, sentences)
    - Overlap to preserve cross-chunk references
    """
    if len(text) <= max_chars:
        return [text]

    chunks = []
    start = 0

    while start < len(text):
        end = start + max_chars

        if end < len(text):
            # Priority 1: Break at section headers (legal docs often have these)
            section_patterns = ["\n[HEADING", "\n[PAGE", "\nARTICLE ", "\nSECTION ", "\nCLAUSE "]
            best_break = -1
            for pattern in section_patterns:
                break_pos = text.rfind(pattern, start + max_chars - overlap, end)
                if break_pos > best_break:
                    best_break = break_pos

            if best_break > start:
                end = best_break
            else:
                # Priority 2: Break at paragraph boundary
                paragraph_break = text.rfind("\n\n", start + max_chars - overlap, end)
                if paragraph_break > start:
                    end = paragraph_break
                else:
                    # Priority 3: Break at sentence
                    sentence_break = text.rfind(". ", start + max_chars - overlap, end)
                    if sentence_break > start:
                        end = sentence_break + 1

        chunks.append(text[start:end].strip())
        start = end - overlap if end < len(text) else end

    return chunks


def truncate_for_context(text: str, max_chars: int = 6000) -> str:
    """
    Truncate text for use as context while preserving key information.
    Used for example documents to fit within token limits.
    """
    if len(text) <= max_chars:
        return text

    # Take beginning and end (most important parts of legal docs)
    beginning = text[:max_chars // 2]
    ending = text[-(max_chars // 2):]

    return f"{beginning}\n\n[... DOCUMENT CONTINUES - {len(text) - max_chars} characters omitted ...]\n\n{ending}"


def estimate_tokens(text: str) -> int:
    """Rough token estimate (1 token ≈ 4 characters for English)."""
    return len(text) // 4


# Global transformer instances
transformer: Optional[SmartDocumentTransformer] = None
multi_transformer: Optional[MultiExampleTransformer] = None
context_aware_transformer: Optional[ContextAwareTransformer] = None


def get_transformer() -> SmartDocumentTransformer:
    global transformer
    if transformer is None:
        configure_dspy()
        transformer = SmartDocumentTransformer()
    return transformer


def get_multi_transformer() -> MultiExampleTransformer:
    global multi_transformer
    if multi_transformer is None:
        configure_dspy()
        multi_transformer = MultiExampleTransformer()
    return multi_transformer


def get_context_aware_transformer() -> ContextAwareTransformer:
    global context_aware_transformer
    if context_aware_transformer is None:
        configure_dspy()
        context_aware_transformer = ContextAwareTransformer()
    return context_aware_transformer


# ============================================================================
# FastAPI Endpoints
# ============================================================================

@app.get("/", response_class=HTMLResponse)
async def serve_frontend():
    """Serve the HTML frontend."""
    html_file = TEMPLATES_DIR / "index.html"
    return html_file.read_text(encoding="utf-8")


@app.post("/transform")
async def transform_document(
    example_input: UploadFile = File(..., description="Example input document (before transformation)"),
    example_output: UploadFile = File(..., description="Example output document (after transformation)"),
    new_document: UploadFile = File(..., description="New document to transform")
):
    """
    Transform a document based on an example input/output pair.

    Accepts PDF, DOCX, TXT, and other text-based formats.
    Files and outputs are saved locally for future reference.
    """
    try:
        # Generate job ID for this transformation
        job_id = generate_job_id()

        # Read file contents
        example_input_bytes = await example_input.read()
        example_output_bytes = await example_output.read()
        new_document_bytes = await new_document.read()

        # Save uploaded files
        save_uploaded_file(example_input_bytes, example_input.filename, job_id, "example_input")
        save_uploaded_file(example_output_bytes, example_output.filename, job_id, "example_output")
        save_uploaded_file(new_document_bytes, new_document.filename, job_id, "new_document")

        # Extract text from all documents
        example_input_text = extract_text(example_input.filename, example_input_bytes)
        example_output_text = extract_text(example_output.filename, example_output_bytes)
        new_document_text = extract_text(new_document.filename, new_document_bytes)

        if not example_input_text.strip():
            raise HTTPException(status_code=400, detail="Example input document is empty")
        if not example_output_text.strip():
            raise HTTPException(status_code=400, detail="Example output document is empty")
        if not new_document_text.strip():
            raise HTTPException(status_code=400, detail="New document is empty")

        # Detect transformation type
        transformation_type = detect_transformation_type(example_input_text, example_output_text)

        # Truncate example documents if too large (to fit in context window)
        # Keep full text for storage but use truncated for API calls
        example_input_for_api = truncate_for_context(example_input_text, max_chars=8000)
        example_output_for_api = truncate_for_context(example_output_text, max_chars=8000)

        # Log document sizes for debugging
        total_example_tokens = estimate_tokens(example_input_for_api) + estimate_tokens(example_output_for_api)

        # Get transformer and run transformation
        trans = get_transformer()

        # Handle large documents with chunking
        chunks = chunk_text(new_document_text)

        if len(chunks) == 1:
            # Single chunk - process directly
            result = trans.forward(
                example_input=example_input_for_api,
                example_output=example_output_for_api,
                new_document=new_document_text
            )
            transformed_text = result.transformed_document
            analysis = result.transformation_analysis
        else:
            # Multiple chunks - process each and combine
            transformed_parts = []
            analyses = []

            for i, chunk in enumerate(chunks):
                # Add chunk context for continuity
                chunk_context = f"[Processing chunk {i+1} of {len(chunks)}]\n\n{chunk}"
                result = trans.forward(
                    example_input=example_input_for_api,
                    example_output=example_output_for_api,
                    new_document=chunk_context
                )
                transformed_parts.append(result.transformed_document)
                if i == 0:  # Only capture analysis from first chunk
                    analyses.append(result.transformation_analysis)

            transformed_text = "\n\n".join(transformed_parts)
            analysis = analyses[0] if analyses else "Multi-chunk transformation"

        # Build metadata
        metadata = {
            "input_length": len(new_document_text),
            "output_length": len(transformed_text),
            "transformation_type": transformation_type,
            "chunks_processed": len(chunks),
            "estimated_tokens_used": total_example_tokens + estimate_tokens(new_document_text),
            "example_truncated": len(example_input_text) > 8000 or len(example_output_text) > 8000,
            "files": {
                "example_input": example_input.filename,
                "example_output": example_output.filename,
                "new_document": new_document.filename
            }
        }

        # Save output and job record
        save_output(job_id, transformed_text, analysis, metadata)
        save_job_record(job_id, {
            "type": "single",
            "metadata": metadata,
            "transformation_analysis": analysis
        })

        return {
            "job_id": job_id,
            "transformation_analysis": analysis,
            "transformed_document": transformed_text,
            "metadata": metadata
        }

    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Transformation failed: {str(e)}")


@app.post("/transform-multi")
async def transform_document_multi(
    example_inputs: list[UploadFile] = File(..., description="List of example input documents"),
    example_outputs: list[UploadFile] = File(..., description="List of example output documents"),
    new_document: UploadFile = File(..., description="New document to transform")
):
    """
    Transform a document using multiple example pairs for better pattern learning.

    Provides improved accuracy by learning from multiple input/output examples.
    Files and outputs are saved locally for future reference.
    """
    try:
        # Generate job ID for this transformation
        job_id = generate_job_id()

        if len(example_inputs) != len(example_outputs):
            raise HTTPException(
                status_code=400,
                detail="Number of example inputs must match number of example outputs"
            )

        if len(example_inputs) < 1:
            raise HTTPException(status_code=400, detail="At least one example pair is required")

        # Extract text from all documents and save files
        examples = []
        example_filenames = []
        for i, (inp, out) in enumerate(zip(example_inputs, example_outputs)):
            inp_bytes = await inp.read()
            out_bytes = await out.read()

            # Save uploaded files
            save_uploaded_file(inp_bytes, inp.filename, job_id, f"example_input_{i+1}")
            save_uploaded_file(out_bytes, out.filename, job_id, f"example_output_{i+1}")

            inp_text = extract_text(inp.filename, inp_bytes)
            out_text = extract_text(out.filename, out_bytes)

            if inp_text.strip() and out_text.strip():
                examples.append({
                    "index": i + 1,
                    "input": inp_text,
                    "output": out_text
                })
                example_filenames.append({
                    "input": inp.filename,
                    "output": out.filename
                })

        if not examples:
            raise HTTPException(status_code=400, detail="No valid example pairs found")

        # Extract and save new document
        new_doc_bytes = await new_document.read()
        save_uploaded_file(new_doc_bytes, new_document.filename, job_id, "new_document")
        new_doc_text = extract_text(new_document.filename, new_doc_bytes)

        if not new_doc_text.strip():
            raise HTTPException(status_code=400, detail="New document is empty")

        # Build examples summary
        examples_summary = "TRANSFORMATION EXAMPLES:\n\n"
        for ex in examples:
            examples_summary += f"=== Example {ex['index']} ===\n"
            examples_summary += f"INPUT:\n{ex['input'][:2000]}{'...' if len(ex['input']) > 2000 else ''}\n\n"
            examples_summary += f"OUTPUT:\n{ex['output'][:2000]}{'...' if len(ex['output']) > 2000 else ''}\n\n"

        # Detect transformation type from first example
        transformation_type = detect_transformation_type(
            examples[0]["input"],
            examples[0]["output"]
        )

        # Get multi-example transformer and run
        trans = get_multi_transformer()

        chunks = chunk_text(new_doc_text)

        if len(chunks) == 1:
            result = trans.forward(
                examples_summary=examples_summary,
                new_document=new_doc_text
            )
            transformed_text = result.transformed_document
            analysis = result.transformation_analysis
        else:
            transformed_parts = []
            analyses = []

            for i, chunk in enumerate(chunks):
                result = trans.forward(
                    examples_summary=examples_summary,
                    new_document=chunk
                )
                transformed_parts.append(result.transformed_document)
                if i == 0:
                    analyses.append(result.transformation_analysis)

            transformed_text = "\n\n".join(transformed_parts)
            analysis = analyses[0] if analyses else "Multi-chunk transformation"

        # Build metadata
        metadata = {
            "input_length": len(new_doc_text),
            "output_length": len(transformed_text),
            "transformation_type": transformation_type,
            "example_pairs_used": len(examples),
            "chunks_processed": len(chunks),
            "files": {
                "example_pairs": example_filenames,
                "new_document": new_document.filename
            }
        }

        # Save output and job record
        save_output(job_id, transformed_text, analysis, metadata)
        save_job_record(job_id, {
            "type": "multi",
            "metadata": metadata,
            "transformation_analysis": analysis
        })

        return {
            "job_id": job_id,
            "transformation_analysis": analysis,
            "transformed_document": transformed_text,
            "metadata": metadata
        }

    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Transformation failed: {str(e)}")


@app.get("/health")
async def health_check():
    """Health check endpoint."""
    return {"status": "healthy", "service": "Smart Document Transformation System"}


@app.get("/jobs")
async def get_jobs():
    """List all saved transformation jobs."""
    jobs = list_jobs()
    return {"jobs": jobs, "total": len(jobs)}


@app.get("/jobs/{job_id}")
async def get_job_details(job_id: str):
    """Get details of a specific transformation job."""
    job = get_job(job_id)
    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    # Load the transformed output
    output_file = OUTPUTS_DIR / job_id / "transformed_output.txt"
    transformed_document = ""
    if output_file.exists():
        transformed_document = output_file.read_text(encoding="utf-8")

    return {
        **job,
        "transformed_document": transformed_document
    }


@app.get("/jobs/{job_id}/download")
async def download_job_output(job_id: str):
    """Download the transformed output of a job."""
    output_file = OUTPUTS_DIR / job_id / "transformed_output.txt"
    if not output_file.exists():
        raise HTTPException(status_code=404, detail="Output file not found")

    return FileResponse(
        path=output_file,
        filename=f"transformed_{job_id}.txt",
        media_type="text/plain"
    )


# ============================================================================
# Background Processing Functions
# ============================================================================

def process_transformation_sync(
    job_id: str,
    example_input_text: str,
    example_output_text: str,
    new_document_text: str,
    transformation_type: str,
    file_metadata: dict
):
    """
    Synchronous transformation processing that runs in a thread pool.
    Uses context-aware chunking for large documents.
    """
    try:
        # Truncate examples for API calls
        example_input_for_api = truncate_for_context(example_input_text, max_chars=8000)
        example_output_for_api = truncate_for_context(example_output_text, max_chars=8000)

        # Chunk the document
        chunks = chunk_text(new_document_text)
        total_chunks = len(chunks)

        update_job_status(
            job_id,
            status=JobStatus.PROCESSING,
            total_chunks=total_chunks,
            message=f"Starting transformation ({total_chunks} chunk{'s' if total_chunks > 1 else ''})"
        )

        if total_chunks == 1:
            # Single chunk - use standard transformer
            update_job_status(
                job_id,
                current_chunk=1,
                progress=10,
                message="Processing document..."
            )

            trans = get_transformer()
            result = trans.forward(
                example_input=example_input_for_api,
                example_output=example_output_for_api,
                new_document=new_document_text
            )

            transformed_text = result.transformed_document
            analysis = result.transformation_analysis

            update_job_status(
                job_id,
                progress=100,
                partial_result=transformed_text,
                analysis=analysis,
                message="Transformation complete"
            )

        else:
            # Multiple chunks - use context-aware transformer
            trans = get_context_aware_transformer()
            transformed_parts = []
            full_analysis = []
            previous_context = ""

            for i, chunk in enumerate(chunks):
                chunk_num = i + 1
                progress = int((chunk_num / total_chunks) * 100)

                update_job_status(
                    job_id,
                    current_chunk=chunk_num,
                    progress=max(5, progress - 5),  # Leave room for completion
                    message=f"Processing chunk {chunk_num} of {total_chunks}..."
                )

                result = trans.forward(
                    example_input=example_input_for_api,
                    example_output=example_output_for_api,
                    previous_context=previous_context,
                    current_chunk=chunk,
                    chunk_info=f"Chunk {chunk_num} of {total_chunks}"
                )

                transformed_parts.append(result.transformed_chunk)
                full_analysis.append(f"[Chunk {chunk_num}] {result.transformation_analysis}")

                # Update context for next chunk
                previous_context = result.context_summary

                # Update partial result for streaming
                current_result = "\n\n".join(transformed_parts)
                update_job_status(
                    job_id,
                    partial_result=current_result,
                    progress=progress
                )

            transformed_text = "\n\n".join(transformed_parts)
            analysis = "\n\n".join(full_analysis)

            update_job_status(
                job_id,
                progress=100,
                partial_result=transformed_text,
                analysis=analysis,
                message="Transformation complete"
            )

        # Build metadata
        metadata = {
            "input_length": len(new_document_text),
            "output_length": len(transformed_text),
            "transformation_type": transformation_type,
            "chunks_processed": total_chunks,
            "estimated_tokens_used": estimate_tokens(example_input_for_api) + estimate_tokens(example_output_for_api) + estimate_tokens(new_document_text),
            "example_truncated": len(example_input_text) > 8000 or len(example_output_text) > 8000,
            "files": file_metadata,
            "processing_mode": "context_aware" if total_chunks > 1 else "standard"
        }

        # Save output and job record
        save_output(job_id, transformed_text, analysis, metadata)
        save_job_record(job_id, {
            "type": "async",
            "metadata": metadata,
            "transformation_analysis": analysis
        })

        update_job_status(
            job_id,
            status=JobStatus.COMPLETED,
            message="Transformation completed successfully"
        )

    except Exception as e:
        update_job_status(
            job_id,
            status=JobStatus.FAILED,
            error=str(e),
            message=f"Transformation failed: {str(e)}"
        )
        # Still save a job record for the failure
        save_job_record(job_id, {
            "type": "async",
            "error": str(e),
            "status": "failed"
        })


async def run_background_transformation(
    job_id: str,
    example_input_text: str,
    example_output_text: str,
    new_document_text: str,
    transformation_type: str,
    file_metadata: dict
):
    """Run transformation in thread pool to avoid blocking the event loop."""
    loop = asyncio.get_event_loop()
    await loop.run_in_executor(
        executor,
        process_transformation_sync,
        job_id,
        example_input_text,
        example_output_text,
        new_document_text,
        transformation_type,
        file_metadata
    )


# ============================================================================
# Async Processing Endpoints
# ============================================================================

@app.post("/transform-async")
async def transform_document_async(
    background_tasks: BackgroundTasks,
    example_input: UploadFile = File(..., description="Example input document (before transformation)"),
    example_output: UploadFile = File(..., description="Example output document (after transformation)"),
    new_document: UploadFile = File(..., description="New document to transform")
):
    """
    Start an async document transformation job.

    Returns immediately with a job_id that can be used to:
    - Poll status via GET /jobs/{job_id}/status
    - Stream updates via GET /jobs/{job_id}/stream

    Ideal for large documents that would timeout with synchronous processing.
    """
    try:
        # Generate job ID
        job_id = generate_job_id()

        # Read file contents
        example_input_bytes = await example_input.read()
        example_output_bytes = await example_output.read()
        new_document_bytes = await new_document.read()

        # Save uploaded files
        save_uploaded_file(example_input_bytes, example_input.filename, job_id, "example_input")
        save_uploaded_file(example_output_bytes, example_output.filename, job_id, "example_output")
        save_uploaded_file(new_document_bytes, new_document.filename, job_id, "new_document")

        # Extract text
        example_input_text = extract_text(example_input.filename, example_input_bytes)
        example_output_text = extract_text(example_output.filename, example_output_bytes)
        new_document_text = extract_text(new_document.filename, new_document_bytes)

        if not example_input_text.strip():
            raise HTTPException(status_code=400, detail="Example input document is empty")
        if not example_output_text.strip():
            raise HTTPException(status_code=400, detail="Example output document is empty")
        if not new_document_text.strip():
            raise HTTPException(status_code=400, detail="New document is empty")

        # Detect transformation type
        transformation_type = detect_transformation_type(example_input_text, example_output_text)

        # Estimate chunks for progress tracking
        chunks = chunk_text(new_document_text)
        total_chunks = len(chunks)

        # Initialize job status
        init_job_status(job_id, total_chunks)

        # File metadata
        file_metadata = {
            "example_input": example_input.filename,
            "example_output": example_output.filename,
            "new_document": new_document.filename
        }

        # Start background processing
        background_tasks.add_task(
            run_background_transformation,
            job_id,
            example_input_text,
            example_output_text,
            new_document_text,
            transformation_type,
            file_metadata
        )

        return {
            "job_id": job_id,
            "status": "pending",
            "message": f"Job queued for processing ({total_chunks} chunk{'s' if total_chunks > 1 else ''})",
            "total_chunks": total_chunks,
            "estimated_document_size": len(new_document_text),
            "links": {
                "status": f"/jobs/{job_id}/status",
                "stream": f"/jobs/{job_id}/stream",
                "result": f"/jobs/{job_id}"
            }
        }

    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e))
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to start transformation: {str(e)}")


@app.get("/jobs/{job_id}/status")
async def get_job_status_endpoint(job_id: str):
    """
    Get the current status of an async transformation job.

    Returns progress information including:
    - Current status (pending, processing, completed, failed)
    - Progress percentage
    - Current chunk being processed
    - Partial results (if available)
    """
    status = get_job_status(job_id)

    if status is None:
        # Check if job exists in storage but not in memory (e.g., after restart)
        job = get_job(job_id)
        if job:
            return {
                "job_id": job_id,
                "status": JobStatus.COMPLETED,
                "progress": 100,
                "message": "Job completed (retrieved from storage)",
                "completed_at": job.get("created_at")
            }
        raise HTTPException(status_code=404, detail="Job not found")

    return status


@app.get("/jobs/{job_id}/stream")
async def stream_job_updates(job_id: str):
    """
    Stream job updates using Server-Sent Events (SSE).

    Provides real-time updates as the transformation progresses.
    Connect to this endpoint to receive progress updates without polling.
    """
    status = get_job_status(job_id)

    if status is None:
        # Check storage
        job = get_job(job_id)
        if job:
            async def completed_stream():
                # Send completed event for jobs retrieved from storage
                output_file = OUTPUTS_DIR / job_id / "transformed_output.txt"
                transformed_document = ""
                if output_file.exists():
                    transformed_document = output_file.read_text(encoding="utf-8")

                event = {
                    "job_id": job_id,
                    "status": JobStatus.COMPLETED,
                    "progress": 100,
                    "message": "Job completed",
                    "transformed_document": transformed_document[:1000] + "..." if len(transformed_document) > 1000 else transformed_document
                }
                yield f"data: {json.dumps(event)}\n\n"

            return StreamingResponse(
                completed_stream(),
                media_type="text/event-stream",
                headers={
                    "Cache-Control": "no-cache",
                    "Connection": "keep-alive",
                    "X-Accel-Buffering": "no"
                }
            )

        raise HTTPException(status_code=404, detail="Job not found")

    async def event_stream():
        last_progress = -1
        last_status = None

        while True:
            current_status = get_job_status(job_id)

            if current_status is None:
                break

            # Send update if progress or status changed
            if current_status["progress"] != last_progress or current_status["status"] != last_status:
                event = {
                    "job_id": job_id,
                    "status": current_status["status"],
                    "progress": current_status["progress"],
                    "current_chunk": current_status["current_chunk"],
                    "total_chunks": current_status["total_chunks"],
                    "message": current_status["message"]
                }

                # Include partial result preview for streaming updates
                if current_status["partial_result"]:
                    preview = current_status["partial_result"]
                    if len(preview) > 500:
                        preview = preview[:500] + "..."
                    event["partial_result_preview"] = preview

                yield f"data: {json.dumps(event)}\n\n"

                last_progress = current_status["progress"]
                last_status = current_status["status"]

            # Check if job is complete
            if current_status["status"] in [JobStatus.COMPLETED, JobStatus.FAILED]:
                # Send final event with full result or error
                final_event = {
                    "job_id": job_id,
                    "status": current_status["status"],
                    "progress": 100,
                    "message": current_status["message"],
                    "final": True
                }

                if current_status["status"] == JobStatus.COMPLETED:
                    final_event["analysis"] = current_status.get("analysis", "")
                else:
                    final_event["error"] = current_status.get("error", "Unknown error")

                yield f"data: {json.dumps(final_event)}\n\n"
                break

            await asyncio.sleep(0.5)  # Poll every 500ms

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no"
        }
    )


# ============================================================================
# Application Entry Point
# ============================================================================

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
